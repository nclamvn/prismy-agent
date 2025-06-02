# === AI Translator Agent: app.py - PHIÊN BẢN CHUYÊN NGHIỆP ===
import streamlit as st
import fitz
import os
import asyncio
import httpx
import tiktoken
import time
import hashlib
import pickle
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import simpleSplit
from difflib import SequenceMatcher
import nest_asyncio
from langdetect import detect
import pandas as pd
import re
from datetime import datetime
from logger_utils import log_model_error
from model_router_async import call_model_async
from api_keys import API_KEYS

nest_asyncio.apply()

# === Font PDF ===
FONT_PATH = "NotoSans-Regular.ttf"
FONT_NAME = "NotoSans"
if os.path.exists(FONT_PATH):
    pdfmetrics.registerFont(TTFont(FONT_NAME, FONT_PATH))
else:
    st.warning("⚠️ Không tìm thấy font NotoSans-Regular.ttf.")

# === DANH SÁCH NGÔN NGỮ PHỔ BIẾN ===
SUPPORTED_LANGUAGES = {
    "Vietnamese": "🇻🇳 Tiếng Việt",
    "English": "🇺🇸 English", 
    "Chinese": "🇨🇳 中文",
    "Japanese": "🇯🇵 日本語",
    "Korean": "🇰🇷 한국어",
    "French": "🇫🇷 Français",
    "German": "🇩🇪 Deutsch",
    "Spanish": "🇪🇸 Español",
    "Thai": "🇹🇭 ภาษาไทย",
    "Arabic": "🇸🇦 العربية"
}

# === Hàm xử lý văn bản và dịch ===
def remove_overlap(prev, curr):
    matcher = SequenceMatcher(None, prev[-200:], curr)
    match = matcher.find_longest_match(0, len(prev[-200:]), 0, len(curr))
    return curr[match.b + match.size:] if match.size > 30 else curr

def count_tokens(text, model="gpt-3.5-turbo"):
    enc = tiktoken.get_encoding("cl100k_base")
    return len(enc.encode(text))

def estimate_cost(text, model):
    tokens = count_tokens(text, model)
    rate = {
        "gpt-3.5-turbo": 0.0035,
        "gpt-4.1-mini": 0.005,
        "gpt-4.1": 0.03,
        "gpt-4": 0.06,
        "gpt-4o": 0.01,
        "gpt-4.5": 0.06
    }
    for k in rate:
        if k in model:
            return tokens, round(tokens / 1000 * rate[k], 4)
    return tokens, round(tokens / 1000 * 0.03, 4)

def split_text(text, max_length=2000, overlap=200):
    words = text.split()
    chunks, i = [], 0
    while i < len(words):
        chunks.append(" ".join(words[i:i + max_length]))
        i += max_length - overlap
    return chunks

def build_prompt(para: str, target_lang: str, style_level: str) -> str:
    prompt = f"Translate this text to {target_lang}:\n\n{para}"
    lvl = style_level.lower()

    if "chuẩn" in lvl and "in ấn" not in lvl:
        prompt += "\n\nStyle: smooth, fluent, professional."
    elif "cao cấp" in lvl or "chuẩn in ấn" in lvl:
        prompt += "\n\nStyle: publication-quality, accurate, natural."
    else:
        prompt += "\n\nStyle: rough, simple draft."

    return prompt

async def translate_paragraphs(
    paragraphs: list[str],
    style_level: str,
    target_lang: str,
    model_name: str
) -> tuple[list[str], list[int]]:
    results: list[str | None] = [None] * len(paragraphs)
    errors: list[int] = []

    async def translate_one(idx: int, para: str) -> None:
        try:
            prompt = build_prompt(para, target_lang, style_level)
            messages = [{"role": "user", "content": prompt}]

            result = await call_model_async(style_level, messages)

            if detect(result[:200]) == detect(para[:200]):
                raise ValueError("Model did not translate.")

            results[idx] = result
        except Exception as exc:
            results[idx] = f"[Lỗi đoạn {idx + 1}]"
            errors.append(idx)
            log_model_error(
                {"model": style_level, "provider": "mixed"},
                f"[APP] Exception đoạn {idx + 1}: {exc}",
                tag="APP"
            )

    await asyncio.gather(*(translate_one(i, p) for i, p in enumerate(paragraphs)))

    results = [
        r if r is not None else f"[Lỗi đoạn {i + 1}]"
        for i, r in enumerate(results)
    ]
    return results, errors

def export_word(paragraphs):
    doc = Document()
    doc.add_heading("Bản dịch tài liệu", level=1)
    for para in paragraphs:
        if '[Lỗi đoạn' in para:
            continue
        p = doc.add_paragraph()
        run = p.add_run(para)
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def export_pdf(paragraphs):
    buf = BytesIO()
    c = canvas.Canvas(buf, pagesize=A4)
    width, height = A4
    margin = 40
    max_width = width - 2 * margin
    y = height - margin
    line_height = 16

    try:
        c.setFont(FONT_NAME, 12)
    except:
        c.setFont("Helvetica", 12)

    for para in paragraphs:
        if '[Lỗi đoạn' in para:
            continue
        lines = simpleSplit(para, FONT_NAME, 12, max_width)
        for line in lines:
            if y < margin + line_height:
                c.showPage()
                y = height - margin
                c.setFont(FONT_NAME, 12)
            c.drawString(margin, y, line)
            y -= line_height
        y -= line_height

    c.save()
    buf.seek(0)
    return buf

def get_file_hash(file_bytes):
    return hashlib.sha256(file_bytes).hexdigest()

def save_translation_cache(
        file_hash: str,
        model: str,
        target_lang: str,
        result: list[str],
        word_buf: BytesIO,
        pdf_buf: BytesIO
) -> None:
    os.makedirs("cache", exist_ok=True)
    cache_key = f"{file_hash}_{model}_{target_lang}"
    with open(f"cache/{cache_key}.pkl", "wb") as f:
        pickle.dump((result, word_buf.getvalue(), pdf_buf.getvalue()), f)

def load_translation_cache(
        file_hash: str,
        model: str,
        target_lang: str
) -> tuple[list[str] | None, BytesIO | None, BytesIO | None]:
    cache_key = f"{file_hash}_{model}_{target_lang}"
    try:
        with open(f"cache/{cache_key}.pkl", "rb") as f:
            result, word_bytes, pdf_bytes = pickle.load(f)
            return result, BytesIO(word_bytes), BytesIO(pdf_bytes)
    except FileNotFoundError:
        return None, None, None

def parse_failed_log():
    rows = []
    try:
        with open("logs/failed_model.log", "r", encoding="utf-8") as f:
            for line in f:
                match = re.match(r"\[(.*?)\] ❌ \[(.*?)\] (\w+) - (.*?): (.*)", line)
                if match:
                    time, tag, provider, model, msg = match.groups()
                    rows.append({
                        "timestamp": time,
                        "tag": tag,
                        "provider": provider,
                        "model": model,
                        "message": msg
                    })
    except FileNotFoundError:
        pass
    return pd.DataFrame(rows)

def parse_benchmark_log():
    rows = []
    try:
        with open("logs/benchmark.log", "r", encoding="utf-8") as f:
            for line in f:
                match = re.match(r"\[(.*?)\] ✅ BENCHMARK (\w+) - (.*?): ([\d.]+)s.*?(\d+) tokens", line)
                if match:
                    time, provider, model, duration, tokens = match.groups()
                    rows.append({
                        "timestamp": time,
                        "provider": provider,
                        "model": model,
                        "duration": float(duration),
                        "tokens": int(tokens)
                    })
    except FileNotFoundError:
        pass
    return pd.DataFrame(rows)

# === GIAO DIỆN CHÍNH - PHIÊN BẢN CHUYÊN NGHIỆP ===
st.set_page_config(
    page_title="AI Translator Agent", 
    layout="wide",
    initial_sidebar_state="collapsed"
)

# === HEADER CHUYÊN NGHIỆP ===
st.markdown("""
<div style='text-align: center; padding: 1rem 0; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); border-radius: 10px; margin-bottom: 2rem;'>
    <h1 style='color: white; margin: 0; font-size: 2.5rem; font-weight: bold;'>
        🌐 AI Translator Agent
    </h1>
    <p style='color: #f0f2f6; margin: 0.5rem 0 0 0; font-size: 1.2rem;'>
        Dịch tài liệu thông minh với công nghệ AI tiên tiến
    </p>
</div>
""", unsafe_allow_html=True)

# === TABS CHÍNH ===
tab1, tab2 = st.tabs(["📘 Dịch tài liệu", "📊 Theo dõi hệ thống"])

# === TAB 1: DỊCH TÀI LIỆU ===
with tab1:
    # === 1. NHẬP API KEY (ĐỨNG ĐẦU) ===
    st.markdown("### 🔑 Cấu hình API")
    api_key = st.text_input(
        "Nhập OpenAI API Key của bạn:", 
        type="password",
        placeholder="sk-...",
        help="API key sẽ được sử dụng để gọi các mô hình GPT"
    )
    if api_key and api_key.strip() and api_key.strip() not in API_KEYS["openai"]:
        API_KEYS["openai"].insert(0, api_key.strip())
        st.session_state["user_key_added"] = True

    st.divider()

    # === 2. CHỌN CHẾ ĐỘ DỊCH & MÔ HÌNH ===
    st.markdown("### 🎯 Cấu hình dịch thuật")
    
    col1, col2 = st.columns(2)
    with col1:
        style = st.selectbox(
            "Chế độ dịch:",
            ["Thô", "Chuẩn", "Chuẩn in ấn"],
            help="• Thô: Nhanh, chi phí thấp\n• Chuẩn: Cân bằng chất lượng-chi phí\n• Chuẩn in ấn: Chất lượng cao nhất"
        )
    
    with col2:
        default_model = {
            "Thô": "gpt-3.5-turbo",
            "Chuẩn": "gpt-4o", 
            "Chuẩn in ấn": "gpt-4.5"
        }[style]
        
        model = st.selectbox(
            "Mô hình AI:",
            ["gpt-3.5-turbo", "gpt-4.1-mini", "gpt-4.1", "gpt-4", "gpt-4o", "gpt-4.5"],
            index=["gpt-3.5-turbo", "gpt-4.1-mini", "gpt-4.1", "gpt-4", "gpt-4o", "gpt-4.5"].index(default_model)
        )

    st.divider()

    # === 3. TẢI FILE ===
    st.markdown("### 📄 Tải tài liệu")
    uploaded_files = st.file_uploader(
        "Chọn file PDF cần dịch:",
        type=["pdf"], 
        accept_multiple_files=True,
        help="Hỗ trợ nhiều file PDF, tối đa 200MB mỗi file"
    )

    if uploaded_files:
        for uploaded_file in uploaded_files:
            st.markdown(f"""
            <div style='background: #f8f9fa; padding: 1rem; border-radius: 8px; margin: 1rem 0; border-left: 4px solid #007bff;'>
                <h4 style='margin: 0; color: #007bff;'>📄 {uploaded_file.name}</h4>
            </div>
            """, unsafe_allow_html=True)
            
            file_bytes = uploaded_file.read()
            file_hash = get_file_hash(file_bytes)

            try:
                doc = fitz.open(stream=file_bytes, filetype="pdf")
                pages_text = [page.get_text("text") for page in doc]
                full_text = "\n\n".join(pages_text)
                doc.close()
            except Exception as e:
                st.error(f"❌ Không thể xử lý file {uploaded_file.name}: {e}")
                continue

            # === 4. PHÁT HIỆN NGÔN NGỮ ===
            st.markdown("#### 🌐 Phát hiện ngôn ngữ")
            try:
                detected_lang = detect(full_text[:500])
                lang_name = {
                    'en': 'English 🇺🇸', 'vi': 'Vietnamese 🇻🇳', 'zh': 'Chinese 🇨🇳',
                    'ja': 'Japanese 🇯🇵', 'ko': 'Korean 🇰🇷', 'fr': 'French 🇫🇷',
                    'de': 'German 🇩🇪', 'es': 'Spanish 🇪🇸', 'th': 'Thai 🇹🇭', 'ar': 'Arabic 🇸🇦'
                }.get(detected_lang, f'{detected_lang.upper()} 🌍')
                
                st.success(f"**Ngôn ngữ phát hiện:** {lang_name}")
            except:
                st.warning("⚠️ Không thể phát hiện ngôn ngữ tự động")

            # === 5. CHỌN NGÔN NGỮ ĐÍCH ===
            st.markdown("#### 🎯 Chọn ngôn ngữ đích")
            target_lang = st.selectbox(
                "Dịch sang:",
                list(SUPPORTED_LANGUAGES.keys()),
                format_func=lambda x: SUPPORTED_LANGUAGES[x],
                key=f"target_{uploaded_file.name}"
            )

            # === 6. ƯỚC LƯỢNG CHI PHÍ ===
            tokens, cost = estimate_cost(full_text, model)
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📊 Tokens", f"{tokens:,}")
            with col2:
                st.metric("💰 Chi phí ước tính", f"${cost:.4f}")
            with col3:
                st.metric("📄 Số trang", f"{len(pages_text)}")

            # === 7. KIỂM TRA CACHE ===
            cached_result, cached_word, cached_pdf = load_translation_cache(
                file_hash, model, target_lang
            )

            if cached_result:
                st.success("✅ Đã có bản dịch trong cache - tiết kiệm chi phí!")
                
                # Hiển thị mẫu từ cache
                st.markdown("#### 👀 Xem trước bản dịch")
                preview_chunks = cached_result[:3] if len(cached_result) >= 3 else cached_result
                
                for i, chunk in enumerate(preview_chunks):
                    if '[Lỗi đoạn' in chunk:
                        continue
                    st.markdown(f"**📝 Đoạn {i+1}:**")
                    st.markdown(f"""
                    <div style='height: 200px; overflow-y: auto; padding: 1rem; background: #f8f9fa; 
                         border: 1px solid #dee2e6; border-radius: 6px; font-family: "Times New Roman", serif;'>
                        {chunk.replace(chr(10), '<br>')}
                    </div>
                    """, unsafe_allow_html=True)
                    st.markdown("---")

                # Nút tải về + Google Drive
                col1, col2, col3, col4 = st.columns(4)
                with col1:
                    st.download_button(
                        "📥 Tải Word",
                        cached_word,
                        file_name=f"{uploaded_file.name}_translated.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                with col2:
                    st.download_button(
                        "📥 Tải PDF", 
                        cached_pdf,
                        file_name=f"{uploaded_file.name}_translated.pdf",
                        mime="application/pdf"
                    )
                with col3:
                    if st.button("☁️ Lưu Google Drive", key=f"gdrive_{uploaded_file.name}"):
                        st.info("🔄 Tính năng Google Drive sẽ sớm ra mắt!")
                with col4:
                    if st.button("🔄 Dịch lại", key=f"retranslate_{uploaded_file.name}"):
                        st.rerun()

                st.info("💡 **Muốn dịch lại?** Thay đổi chế độ dịch hoặc mô hình AI ở trên rồi nhấn 'Dịch lại'")

            else:
                # === 8. NÚT DỊCH ===
                if st.button(
                    f"🚀 Bắt đầu dịch {uploaded_file.name}",
                    type="primary",
                    use_container_width=True
                ):
                    # Thanh tiến trình
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    with st.spinner("🔄 Đang xử lý tài liệu..."):
                        start_time = time.time()
                        
                        # Cập nhật tiến trình
                        progress_bar.progress(20)
                        status_text.text("📝 Đang chia nhỏ văn bản...")
                        
                        paragraphs = split_text(full_text)
                        
                        progress_bar.progress(40)
                        status_text.text("🤖 Đang dịch với AI...")
                        
                        # FIX: Dùng asyncio.run
                        translated, errors = asyncio.get_event_loop().run_until_complete(
                            translate_paragraphs(paragraphs, style, target_lang, model)
                        )

                        if not translated or not isinstance(translated, list) or not translated[0]:
                            st.error("❌ Không có kết quả dịch hợp lệ.")
                            continue

                        progress_bar.progress(70)
                        status_text.text("🔧 Đang tối ưu kết quả...")
                        
                        # Xử lý overlap
                        final_result = [translated[0]]
                        for i in range(1, len(translated)):
                            final_result.append(remove_overlap(final_result[-1], translated[i]))

                        progress_bar.progress(90)
                        status_text.text("📄 Đang tạo file...")
                        
                        word_buf = export_word(final_result)
                        pdf_buf = export_pdf(final_result)
                        save_translation_cache(
                            file_hash, model, target_lang, final_result, word_buf, pdf_buf
                        )

                        progress_bar.progress(100)
                        duration = round(time.time() - start_time, 1)
                        status_text.text(f"✅ Hoàn thành trong {duration} giây!")

                    # === 9. HIỂN THỊ KẾT QUẢ ===
                    st.success(f"🎉 **Dịch hoàn tất!** Thời gian: **{duration} giây**")
                    
                    if errors:
                        st.warning(f"⚠️ Có {len(errors)} đoạn bị lỗi trong quá trình dịch")

                    # === 10. CỬA SỔ XEM THỬ ===
                    st.markdown("#### 👀 Xem trước kết quả dịch")
                    
                    # Chọn số đoạn hiển thị thông minh
                    total_chunks = len(final_result)
                    if total_chunks <= 1:
                        preview_chunks = final_result
                        st.info("📄 Tài liệu ngắn - hiển thị toàn bộ")
                    elif total_chunks <= 3:
                        preview_chunks = final_result
                        st.info(f"📄 Hiển thị {total_chunks} đoạn")
                    else:
                        preview_chunks = final_result[:3]
                        st.info(f"📄 Hiển thị 3/{total_chunks} đoạn đầu tiên")

                    for i, para in enumerate(preview_chunks):
                        if '[Lỗi đoạn' in para:
                            continue
                        st.markdown(f"**📝 Đoạn {i+1}:**")
                        st.markdown(f"""
                        <div style='height: 300px; overflow-y: auto; padding: 1rem; background: #f8f9fa; 
                             border: 1px solid #dee2e6; border-radius: 6px; font-family: "Times New Roman", serif;
                             line-height: 1.6;'>
                            {para.replace(chr(10), '<br>')}
                        </div>
                        """, unsafe_allow_html=True)
                        st.markdown("---")

                    # === 11. NÚT TẢI VỀ & GOOGLE DRIVE ===
                    st.markdown("#### 💾 Lưu trữ & Tải về")
                    
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.download_button(
                            "📥 Tải file Word",
                            word_buf,
                            file_name=f"{uploaded_file.name}_translated.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                    with col2:
                        st.download_button(
                            "📥 Tải file PDF",
                            pdf_buf,
                            file_name=f"{uploaded_file.name}_translated.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                    with col3:
                        if st.button("☁️ Lưu Google Drive", use_container_width=True):
                            st.info("🔄 Tính năng Google Drive đang phát triển!")

                    # === 12. THÔNG BÁO DỊCH LẠI ===
                    st.markdown("""
                    <div style='background: #e3f2fd; padding: 1rem; border-radius: 8px; border-left: 4px solid #2196f3; margin-top: 1rem;'>
                        <strong>💡 Không hài lòng với kết quả?</strong><br>
                        Bạn có thể thay đổi <strong>chế độ dịch</strong> hoặc <strong>mô hình AI</strong> ở phía trên, 
                        sau đó nhấn <strong>"Dịch lại"</strong> để có kết quả tốt hơn.
                    </div>
                    """, unsafe_allow_html=True)

# === TAB 2: DASHBOARD (GIỮ NGUYÊN) ===
with tab2:
    st.subheader("📈 Hiệu suất mô hình")
    df_bench = parse_benchmark_log()
    if df_bench.empty:
        st.warning("Chưa có dữ liệu benchmark.")
    else:
        avg = df_bench.groupby(["provider", "model"]).agg({"duration": "mean", "tokens": "mean"}).round(2)
        st.dataframe(avg, use_container_width=True)

        count_df = df_bench.groupby(["provider", "model"]).size().reset_index(name="calls")
        st.bar_chart(count_df.set_index("model")["calls"])

        st.caption("🧾 Các bản ghi benchmark gần nhất")
        st.dataframe(df_bench.sort_values("timestamp", ascending=False).head(10), use_container_width=True)

    st.subheader("⚠️ Thống kê lỗi API")
    df_fail = parse_failed_log()
    if df_fail.empty:
        st.success("Không có lỗi nào được ghi nhận.")
    else:
        error_stats = df_fail.groupby(["provider", "model", "tag"]).size().reset_index(name="errors")
        st.dataframe(error_stats, use_container_width=True)

        st.caption("❌ Các lỗi gần nhất")
        st.dataframe(df_fail.sort_values("timestamp", ascending=False).head(10), use_container_width=True)