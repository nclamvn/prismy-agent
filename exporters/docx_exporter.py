from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from typing import Dict, Any, Optional
import logging

logger = logging.getLogger(__name__)

class DOCXExporter:
    """Export processed content to DOCX format with formatting preservation"""
    
    def __init__(self):
        self.document = Document()
        self._setup_default_styles()
    
    def _setup_default_styles(self):
        """Setup default document styles"""
        style = self.document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
    
    def add_heading(self, text: str, level: int = 1):
        """Add a heading with specified level"""
        self.document.add_heading(text, level=level)
    
    def add_paragraph(self, text: str, alignment: str = 'LEFT'):
        """Add a paragraph with specified alignment"""
        paragraph = self.document.add_paragraph(text)
        if alignment.upper() == 'CENTER':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment.upper() == 'RIGHT':
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    def add_table(self, data: pd.DataFrame, include_header: bool = True):
        """Add a table from pandas DataFrame"""
        try:
            rows, cols = data.shape
            table = self.document.add_table(rows=rows + (1 if include_header else 0), cols=cols)
            table.style = 'Table Grid'
            
            # Add headers
            if include_header:
                header_cells = table.rows[0].cells
                for i, column in enumerate(data.columns):
                    header_cells[i].text = str(column)
            
            # Add data
            for i in range(rows):
                row_cells = table.rows[i + (1 if include_header else 0)].cells
                for j in range(cols):
                    row_cells[j].text = str(data.iloc[i, j])
                    
        except Exception as e:
            logger.error(f"Failed to add table: {e}")
            self.add_paragraph("[Table conversion failed]")
    
    def add_image(self, image_path: str, width: Optional[float] = None):
        """Add an image with optional width specification"""
        try:
            if width:
                self.document.add_picture(image_path, width=Inches(width))
            else:
                self.document.add_picture(image_path)
        except Exception as e:
            logger.error(f"Failed to add image {image_path}: {e}")
            self.add_paragraph(f"[Image insertion failed: {image_path}]")
    
    def export_content(self, content: Dict[str, Any], output_path: str):
        """Export processed content to DOCX file"""
        try:
            # Add title if available
            if 'title' in content:
                self.add_heading(content['title'])
            
            # Add main content
            if 'text' in content:
                # Handle text content that might be a string or dict
                text_content = content['text']
                if isinstance(text_content, dict):
                    text_content = str(text_content)
                self.add_paragraph(text_content)
            
            # Add tables if available
            if 'tables' in content and content['tables']:
                self.add_heading('Tables', level=2)
                for table in content['tables']:
                    if isinstance(table, pd.DataFrame):
                        self.add_table(table)
                    elif isinstance(table, dict) and 'dataframe' in table:
                        self.add_table(table['dataframe'])
            
            # Add translation if available
            if 'translation' in content:
                self.add_heading('Translation', level=2)
                if isinstance(content['translation'], dict):
                    translation_text = content['translation'].get('translated_text', '')
                else:
                    translation_text = str(content['translation'])
                self.add_paragraph(translation_text)
            
            # Save document
            self.document.save(output_path)
            logger.info(f"Successfully exported DOCX to {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to export DOCX: {e}")
            return False
    
    def export_translation_result(self, original_text: str, translated_text: str, 
                                metadata: Dict[str, Any], output_path: str):
        """Export translation result with original and translated text"""
        try:
            # Add title
            self.add_heading('Translation Result')
            
            # Add metadata
            self.add_heading('Processing Information', level=2)
            metadata_text = (
                f"Source Language: {metadata.get('detected_lang', 'Unknown').upper()}\n"
                f"Target Language: {metadata.get('target_lang', 'Unknown')}\n"
                f"Processing Time: {metadata.get('processing_time', 0):.2f}s\n"
                f"Confidence Score: {metadata.get('confidence', 0)*100:.1f}%"
            )
            self.add_paragraph(metadata_text)
            
            # Add original text
            self.add_heading('Original Text', level=2)
            self.add_paragraph(original_text)
            
            # Add translated text
            self.add_heading('Translation', level=2)
            self.add_paragraph(translated_text)
            
            # Save document
            self.document.save(output_path)
            logger.info(f"Successfully exported translation to {output_path}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to export translation: {e}")
            return False
